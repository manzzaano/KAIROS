"""
Añade cabecera y pie de página a Entrega1_KAIROS.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX = Path(r"C:\Users\Ismael\Desktop\KAIROS\entregas\entrega1\Entrega1_KAIROS.docx")


def insert_page_field(run):
    """Inserta campo {PAGE} en el run."""
    for tag, text in [('begin', None), ('instrText', 'PAGE'), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.text = text
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        run._r.append(el)


def set_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()

    # Left: project name
    run_l = p.add_run("KAIROS · Entrega 1 – Base funcional real")
    run_l.font.size = Pt(9)
    run_l.font.color.rgb = RGBColor(80, 80, 80)

    # Tab to right
    p.add_run("\t")

    # Right: student name
    run_r = p.add_run("Ismael Manzano León")
    run_r.font.size = Pt(9)
    run_r.font.color.rgb = RGBColor(80, 80, 80)
    run_r.bold = True

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tab stop at right margin for right-aligned text
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9070')   # ~16cm in twips (1cm = 567 twips)
    tabs.append(tab)
    pPr.append(tabs)

    # Horizontal rule below header
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("— ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    run_num = p.add_run()
    run_num.font.size = Pt(9)
    run_num.font.color.rgb = RGBColor(100, 100, 100)
    insert_page_field(run_num)

    run2 = p.add_run(" —")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(100, 100, 100)


doc = Document(str(DOCX))

for section in doc.sections:
    section.different_first_page_header_footer = True  # portada sin cabecera/pie
    set_header(section)
    set_footer(section)

doc.save(str(DOCX))
print(f"Patched: {DOCX}")
