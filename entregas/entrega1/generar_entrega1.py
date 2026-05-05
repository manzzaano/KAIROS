"""
Genera el PDF de Entrega 1 - KAIROS App
Requiere: python-docx, Pillow
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCREENS_DIR = Path(r"C:\Users\Ismael\Desktop\KAIROS\entregas\entrega1\screens")
OUTPUT = Path(r"C:\Users\Ismael\Desktop\KAIROS\entregas\entrega1\Entrega1_KAIROS.docx")

# ──────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def add_code_block(doc, code, lang="dart"):
    style_name = "Code Block"
    if style_name not in [s.name for s in doc.styles]:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(9)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Cm(0.5)

    for line in code.split("\n"):
        p = doc.add_paragraph(style=style_name)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_screenshot(doc, filename, caption=""):
    path = SCREENS_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(3.2))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
    else:
        add_para(doc, f"[Captura no disponible: {filename}]", italic=True, color=(180, 0, 0))

def add_toc(doc):
    """Inserta campo TOC de Word con hipervínculos. Requiere Ctrl+A → F9 en Word para actualizar."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_end)


# ── Cover page helpers ────────────────────────────────────

def _cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _row_height(row, twips):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(twips))
    h.set(qn('w:hRule'), 'exact')
    trPr.append(h)

def _no_borders(tbl):
    tblPr = tbl._tbl.get_or_add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def _cell_no_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _cpara(cell, text='', size=11, bold=False, italic=False,
           color=None, align=WD_ALIGN_PARAGRAPH.CENTER,
           sb=0, sa=0, font='Calibri'):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        r.font.name  = font
        r.font.size  = Pt(size)
        r.bold       = bold
        r.italic     = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p

def _insert_page_field(run):
    for tag in ('begin',):
        fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), tag)
        run._r.append(fc)
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    run._r.append(instr)
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run._r.append(fc2)

def build_portada(doc):
    """Portada moderna de 3 filas: fondo oscuro | franja naranja | sección blanca."""
    # Full content width: 21cm - 2*2.5cm = 16cm = 9072 twips
    tbl = doc.add_table(rows=3, cols=1)
    tbl.style = 'Table Grid'
    _no_borders(tbl)

    tblPr = tbl._tbl.get_or_add_tblPr()
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9072')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Heights (A4 content ~24.7cm = 13985 twips)
    _row_height(tbl.rows[0], 6000)   # dark   ~10.6 cm
    _row_height(tbl.rows[1], 220)    # stripe ~0.39 cm
    _row_height(tbl.rows[2], 7765)   # white  ~13.7 cm

    # ── Row 0: dark ──────────────────────────────────────
    dc = tbl.rows[0].cells[0]
    _cell_shading(dc, '0A0A0A')
    _cell_no_borders(dc)
    dc.paragraphs[0].clear()

    _cpara(dc, sb=50, sa=0)          # top spacer
    _cpara(dc, sb=18, sa=0)
    _cpara(dc, sb=18, sa=0)

    # KAIROS title
    p_k = dc.add_paragraph()
    p_k.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_k.paragraph_format.space_before = Pt(6)
    p_k.paragraph_format.space_after  = Pt(0)
    rk = p_k.add_run("KAIROS")
    rk.font.name = 'Calibri'
    rk.font.size = Pt(64)
    rk.bold = True
    rk.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Orange rule below title
    p_rule = dc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule.paragraph_format.space_before = Pt(6)
    p_rule.paragraph_format.space_after  = Pt(10)
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '18')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'FB923C')
    pBdr.append(bot)
    pPr.append(pBdr)

    _cpara(dc, 'Aplicación de gestión de tareas y productividad personal',
           size=12, color=(0xCC, 0xCC, 0xCC), sb=4, sa=6)
    _cpara(dc, 'Una tarea. Un cronómetro. Sin distracciones.',
           size=10, italic=True, color=(0xFB, 0x92, 0x3C), sb=0, sa=0)

    # ── Row 1: orange stripe ──────────────────────────────
    oc = tbl.rows[1].cells[0]
    _cell_shading(oc, 'FB923C')
    _cell_no_borders(oc)
    oc.paragraphs[0].clear()

    # ── Row 2: white info ─────────────────────────────────
    wc = tbl.rows[2].cells[0]
    _cell_shading(wc, 'FFFFFF')
    _cell_no_borders(wc)
    wc.paragraphs[0].clear()

    _cpara(wc, sb=14, sa=0)

    # "ENTREGA 1" badge
    _cpara(wc, 'ENTREGA 1', size=24, bold=True,
           color=(0xFB, 0x92, 0x3C), sb=6, sa=2)
    _cpara(wc, 'Base funcional real  ·  20%', size=11,
           color=(0x80, 0x80, 0x80), sb=0, sa=10)

    # Horizontal rule
    p_hr = wc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(0)
    p_hr.paragraph_format.space_after  = Pt(8)
    hrPr = p_hr._p.get_or_add_pPr()
    hrBdr = OxmlElement('w:pBdr')
    hrBot = OxmlElement('w:bottom')
    hrBot.set(qn('w:val'), 'single')
    hrBot.set(qn('w:sz'), '4')
    hrBot.set(qn('w:space'), '1')
    hrBot.set(qn('w:color'), 'E0E0E0')
    hrBdr.append(hrBot)
    hrPr.append(hrBdr)

    # Info table (nested, 2 cols, centered)
    info_data = [
        ("Alumno",               "Ismael Manzano León"),
        ("Ciclo Formativo",      "[Nombre Ciclo Formativo]"),
        ("Centro",               "[Nombre del Centro Educativo]"),
        ("Módulo / Asignatura",  "[Nombre del Módulo]"),
        ("Profesor/a",           "[Nombre del Profesor/a]"),
        ("Fecha de entrega",     "27 de abril de 2026"),
        ("Versión",              "1.0.0"),
    ]
    itbl = wc.add_table(rows=len(info_data), cols=2)
    itbl.style = 'Table Grid'
    _no_borders(itbl)

    itblPr = itbl._tbl.get_or_add_tblPr()
    itblW = OxmlElement('w:tblW')
    itblW.set(qn('w:w'), '7200')
    itblW.set(qn('w:type'), 'dxa')
    itblPr.append(itblW)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    itblPr.append(jc)

    for i, (label, value) in enumerate(info_data):
        lc = itbl.rows[i].cells[0]
        vc = itbl.rows[i].cells[1]
        bg = 'F5F5F5' if i % 2 == 0 else 'FAFAFA'
        _cell_shading(lc, bg)
        _cell_shading(vc, bg)
        _cell_no_borders(lc)
        _cell_no_borders(vc)

        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(4)
        lr = lp.add_run(label)
        lr.font.name = 'Calibri'
        lr.font.size = Pt(9)
        lr.bold = True
        lr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vp.paragraph_format.space_before = Pt(4)
        vp.paragraph_format.space_after  = Pt(4)
        vr = vp.add_run(f"  {value}")
        vr.font.name = 'Calibri'
        vr.font.size = Pt(9)
        vr.bold = (i == 0)
        vr.font.color.rgb = RGBColor(0x15, 0x15, 0x15)

    _cpara(wc, sb=10, sa=0)
    _cpara(wc, 'KAIROS 2.0.1  ·  BUILD 2026.04.27  ·  ©IML',
           size=7, color=(0xBB, 0xBB, 0xBB), sb=6, sa=0, font='Courier New')


def apply_headers_footers(doc):
    """Cabecera con nombre del proyecto + alumno. Pie con número de página."""
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # Header
        hdr = section.header
        hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        rl = hp.add_run("KAIROS · Entrega 1 – Base funcional real")
        rl.font.size = Pt(9)
        rl.font.name = 'Calibri'
        rl.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        hp.add_run("\t")

        rr = hp.add_run("Ismael Manzano León")
        rr.font.size = Pt(9)
        rr.font.name = 'Calibri'
        rr.bold = True
        rr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        # Right-aligned tab stop at right margin (~9 cm = 5103 twips from left)
        pPr = hp._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9072')
        tabs.append(tab)
        pPr.append(tabs)

        # Thin bottom rule on header
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'DDDDDD')
        pBdr.append(bot)
        pPr.append(pBdr)

        # Footer
        ftr = section.footer
        ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for txt in ('— ', None, ' —'):
            r = fp.add_run(txt or '')
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
            if txt is None:
                _insert_page_field(r)


def add_two_screenshots(doc, file1, cap1, file2, cap2):
    """Intenta poner dos capturas lado a lado en una tabla."""
    p1 = SCREENS_DIR / file1
    p2 = SCREENS_DIR / file2
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    # fila imágenes
    c1 = table.cell(0, 0)
    c2 = table.cell(0, 1)
    for cell, path in ((c1, p1), (c2, p2)):
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            run = para.add_run()
            run.add_picture(str(path), width=Inches(2.8))
        else:
            para.add_run(f"[{path.name} no encontrado]")
    # fila captions
    for cell, cap in ((table.cell(1, 0), cap1), (table.cell(1, 1), cap2)):
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(cap)
        run.italic = True
        run.font.size = Pt(9)
    doc.add_paragraph()


# ──────────────────────────────────────────────────────────
# DOCUMENT
# ──────────────────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── PORTADA ──────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("KAIROS")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run("Aplicación de Gestión de Tareas y Productividad")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub2.add_run("Entrega 1 – Base funcional real (20%)")
r3.bold = True
r3.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

info_data = [
    ("Alumno",          "Ismael Manzano León"),
    ("Ciclo Formativo", "[Nombre Ciclo Formativo]"),
    ("Centro",          "[Nombre del Centro Educativo]"),
    ("Módulo / Asignatura", "[Nombre del Módulo]"),
    ("Profesor/a",      "[Nombre del Profesor/a]"),
    ("Fecha de entrega","27 de abril de 2026"),
    ("Versión",         "1.0.0"),
]

for label, value in info_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l = p.add_run(f"{label}: ")
    r_l.bold = True
    r_l.font.size = Pt(12)
    r_v = p.add_run(value)
    r_v.font.size = Pt(12)

doc.add_page_break()

# ── ÍNDICE ───────────────────────────────────────────────
add_heading(doc, "Índice")
add_toc(doc)
p_note = doc.add_paragraph()
r_note = p_note.add_run("Nota: abre en Word y pulsa Ctrl+A → F9 para actualizar el índice.")
r_note.italic = True
r_note.font.size = Pt(9)
r_note.font.color.rgb = RGBColor(150, 100, 0)
doc.add_page_break()

# ── 1. DESCRIPCIÓN DEL PROYECTO ──────────────────────────
add_heading(doc, "1. Descripción del proyecto")
add_para(doc,
    "KAIROS es una aplicación móvil de gestión de tareas y productividad personal desarrollada en Flutter. "
    "Implementa Clean Architecture (Domain / Data / Presentation) con BLoC/Cubit como gestor de estado, "
    "Realm como base de datos local offline-first y GoRouter para la navegación. "
    "El proyecto está orientado a ofrecer una experiencia fluida, personalizable y sin dependencia de red.",
    size=11)
doc.add_paragraph()
add_para(doc, "Tecnologías principales:", bold=True, size=11)
tecnologias = [
    "Flutter 3.x (Dart 3)",
    "flutter_bloc / Cubit – gestión de estado reactivo",
    "Realm – base de datos local embebida (offline-first)",
    "go_router – navegación declarativa",
    "get_it – inyección de dependencias",
    "shared_preferences – persistencia de preferencias de usuario",
]
for t in tecnologias:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(t).font.size = Pt(11)

doc.add_paragraph()

# ── 2. ARQUITECTURA ──────────────────────────────────────
add_heading(doc, "2. Arquitectura del proyecto")
add_para(doc,
    "El proyecto sigue Clean Architecture dividido en tres capas claramente separadas:",
    size=11)
capas = [
    ("Domain layer", "Entidades puras (Task, Priority, EnergyLevel) y contratos de repositorio (ITaskRepository). Sin dependencias de Flutter ni Realm."),
    ("Data layer", "Implementación de repositorios (TaskRepositoryImpl), datasource de Realm (TaskRealmDataSource) y modelos de objeto Realm (TaskObject)."),
    ("Presentation layer", "BLoC (TaskBloc, FocusBloc) y Cubits (ThemeCubit) que emiten estados. Widgets reactivos que escuchan esos estados mediante BlocBuilder."),
]
for nombre, desc in capas:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{nombre}: ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

doc.add_paragraph()
add_para(doc, "Estructura de directorios relevante:", bold=True, size=11)
add_code_block(doc, """\
lib/
├── app.dart                          # MaterialApp con ThemeCubit
├── main.dart                         # Punto de entrada + init DI
├── core/
│   ├── di/injection_container.dart   # GetIt – registro de dependencias
│   ├── router/app_router.dart        # GoRouter – rutas declarativas
│   └── theme/
│       ├── kairos_colors.dart        # ThemeExtension con 15 colores dinámicos
│       ├── theme_cubit.dart          # Cubit para modo oscuro y color de acento
│       └── app_theme.dart            # ThemeData dark/light dinámico
└── features/
    ├── tasks/                        # CRUD de tareas (domain, data, presentation)
    ├── dashboard/                    # Pantalla principal
    ├── focus/                        # Modo enfoque + Pomodoro
    ├── stats/                        # Estadísticas calculadas desde Realm
    └── profile/                      # Ajustes de apariencia""", "text")

doc.add_page_break()

# ── 3. FUNCIONALIDAD IMPLEMENTADA ────────────────────────
add_heading(doc, "3. Funcionalidades implementadas")

funcionalidades = [
    ("Gestión de tareas (Realm)",
     "Creación, edición, eliminación y marcado como completada de tareas almacenadas en Realm. "
     "Cada tarea incluye título, descripción, prioridad (alta/media/baja), nivel de energía (1-5), "
     "tiempo estimado, fecha de vencimiento, timestamps de creación y completado."),
    ("Dashboard",
     "Pantalla principal con saludo dinámico, barra de energía acumulada del día, "
     "listado de tareas pendientes con acceso directo a detalle."),
    ("Lista de tareas",
     "Vista completa con filtros (Todas / Pendientes / Completadas), toggle de estado, "
     "acceso a detalle y botón de eliminar."),
    ("Detalle de tarea",
     "Vista de lectura con todos los campos, opción de marcar como completada/pendiente y eliminar."),
    ("Creación de tarea",
     "Formulario completo: título, descripción opcional, selector de prioridad, selector de energía (E1-E5), "
     "estimación de tiempo y fecha límite con DatePicker nativo."),
    ("Modo Enfoque (Pomodoro)",
     "Temporizador Pomodoro de 25 minutos con CustomPainter de arco animado. "
     "Controles: play/pausa, reset, salir. Sesión libre o vinculada a una tarea."),
    ("Estadísticas reales",
     "KPIs calculados sobre datos reales de Realm: tareas completadas, tiempo estimado acumulado, "
     "racha de días consecutivos, promedio diario. Gráfico de barras últimos 7 días y heatmap 28 días."),
    ("Tema dinámico (oscuro/claro + acento)",
     "ThemeCubit con persistencia en SharedPreferences. Paleta de 8 colores de acento seleccionable. "
     "Toggle dark/light en pantalla de Perfil. Cambio instantáneo en toda la aplicación."),
    ("Navegación real (GoRouter)",
     "Navegación por 5 pestañas principales (Hoy / Tareas / Enfoque / Stats / Perfil) mediante "
     "StatefulShellRoute. Rutas anidadas para detalle de tarea, creación y temporizador Pomodoro."),
]

for nombre, desc in funcionalidades:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(f"{nombre}: ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

doc.add_page_break()

# ── 4. FRAGMENTOS DE CÓDIGO RELEVANTES ───────────────────
add_heading(doc, "4. Fragmentos de código relevantes")

# 4.1 Entidad Task
add_heading(doc, "4.1 Entidad de dominio – Task", level=2)
add_para(doc,
    "Entidad pura de dominio que representa una tarea. Sin dependencias de frameworks externos.",
    size=11)
add_code_block(doc, """\
// lib/features/tasks/domain/entities/task.dart
enum Priority { high, medium, low }

class Task {
  final String id;
  final String title;
  final String description;
  final Priority priority;
  final int energyLevel;       // 1-5
  final int estimateMinutes;
  final bool isDone;
  final DateTime? dueDate;
  final DateTime? completedAt; // timestamp real de completado
  final DateTime? createdAt;

  bool get isOverdue =>
      dueDate != null && !isDone && dueDate!.isBefore(DateTime.now());
}""")

# 4.2 Realm datasource
add_heading(doc, "4.2 Realm datasource – TaskRealmDataSource", level=2)
add_para(doc,
    "Fuente de datos local que persiste y consulta tareas en Realm. "
    "Se eliminó cualquier semilla de datos simulados — todos los datos son reales.",
    size=11)
add_code_block(doc, """\
// lib/features/tasks/data/datasources/task_realm_datasource.dart
class TaskRealmDataSource {
  final Realm realm;
  TaskRealmDataSource(this.realm);

  List<Task> getTasks() =>
      realm.all<TaskObject>().map(_toEntity).toList();

  void createTask(Task task) {
    realm.write(() {
      realm.add(TaskObject(
        task.id, task.title, task.description,
        task.priority.name, task.energyLevel,
        task.estimateMinutes, false,
        dueDate: task.dueDate?.toUtc(),
        createdAt: DateTime.now().toUtc(),
      ));
    });
  }

  void toggleTask(String id) {
    final obj = realm.find<TaskObject>(id);
    if (obj == null) return;
    realm.write(() {
      obj.isDone = !obj.isDone;
      obj.completedAt = obj.isDone ? DateTime.now().toUtc() : null;
    });
  }

  void deleteTask(String id) {
    final obj = realm.find<TaskObject>(id);
    if (obj != null) realm.write(() => realm.delete(obj));
  }
}""")

# 4.3 ThemeCubit
add_heading(doc, "4.3 Sistema de tema dinámico – ThemeCubit", level=2)
add_para(doc,
    "Cubit que gestiona el modo oscuro/claro y el color de acento, "
    "persiste ambas preferencias en SharedPreferences.",
    size=11)
add_code_block(doc, """\
// lib/core/theme/theme_cubit.dart
class ThemeState {
  final ThemeMode mode;
  final Color accent;
  const ThemeState({required this.mode, required this.accent});
}

class ThemeCubit extends Cubit<ThemeState> {
  ThemeCubit()
      : super(const ThemeState(
            mode: ThemeMode.dark,
            accent: Color(0xFFFB923C)));

  Future<void> load() async {
    final prefs = await SharedPreferences.getInstance();
    final isDark = prefs.getBool('isDark') ?? true;
    final accentVal = prefs.getInt('accent') ?? 0xFFFB923C;
    emit(ThemeState(
      mode: isDark ? ThemeMode.dark : ThemeMode.light,
      accent: Color(accentVal),
    ));
  }

  Future<void> toggleMode() async {
    final isDark = state.mode == ThemeMode.light;
    final prefs = await SharedPreferences.getInstance();
    await prefs.setBool('isDark', isDark);
    emit(ThemeState(mode: isDark ? ThemeMode.dark : ThemeMode.light,
                    accent: state.accent));
  }

  Future<void> setAccent(Color color) async {
    final prefs = await SharedPreferences.getInstance();
    await prefs.setInt('accent', color.value);
    emit(ThemeState(mode: state.mode, accent: color));
  }
}""")

# 4.4 Stats
add_heading(doc, "4.4 Estadísticas reales desde Realm – StatsPage", level=2)
add_para(doc,
    "Las estadísticas se calculan en tiempo real sobre los datos de Realm. "
    "No existe ningún dato simulado ni hardcodeado.",
    size=11)
add_code_block(doc, """\
// lib/features/stats/presentation/pages/stats_page.dart (fragmento)
static List<int> _weekBars(List<Task> tasks) {
  final today = _dayOnly(DateTime.now());
  return List.generate(7, (i) {
    final day = today.subtract(Duration(days: 6 - i));
    return tasks
        .where((t) => t.completedAt != null &&
                      _dayOnly(t.completedAt!) == day)
        .length;
  });
}

static int _streak(List<Task> tasks) {
  int count = 0;
  DateTime check = _dayOnly(DateTime.now());
  while (true) {
    final hasAny = tasks.any((t) =>
        t.completedAt != null && _dayOnly(t.completedAt!) == check);
    if (!hasAny) break;
    count++;
    check = check.subtract(const Duration(days: 1));
  }
  return count;
}""")

# 4.5 Router
add_heading(doc, "4.5 Navegación declarativa – GoRouter", level=2)
add_para(doc,
    "Navegación estructurada con StatefulShellRoute para mantener el estado de cada pestaña "
    "y rutas anidadas para pantallas de detalle.",
    size=11)
add_code_block(doc, """\
// lib/core/router/app_router.dart (fragmento)
final router = GoRouter(
  initialLocation: '/dashboard',
  routes: [
    StatefulShellRoute.indexedStack(
      builder: (_, __, shell) => AppShell(shell: shell),
      branches: [
        StatefulShellBranch(routes: [GoRoute(path: '/dashboard', ...)]),
        StatefulShellBranch(routes: [
          GoRoute(path: '/tasks',
            routes: [
              GoRoute(path: 'new', ...),
              GoRoute(path: ':id', ...),
            ]),
        ]),
        StatefulShellBranch(routes: [
          GoRoute(path: '/focus',
            routes: [GoRoute(path: 'timer', ...)]),
        ]),
        StatefulShellBranch(routes: [GoRoute(path: '/stats', ...)]),
        StatefulShellBranch(routes: [GoRoute(path: '/profile', ...)]),
      ],
    ),
  ],
);""")

doc.add_page_break()

# ── 5. CAPTURAS DE PANTALLA ───────────────────────────────
add_heading(doc, "5. Capturas de pantalla – Aplicación en ejecución real")
add_para(doc,
    "Las siguientes capturas muestran la aplicación KAIROS ejecutándose en un emulador Android "
    "(Medium Phone API 35 – x86_64) con el SDK de Flutter compilado en modo debug.",
    size=11)
doc.add_paragraph()

screens = [
    ("01_dashboard_empty.png", "Dashboard (Hoy) – pantalla principal con 0 tareas"),
    ("02_create_task.png",     "Crear tarea – formulario con todos los campos"),
    ("03_dashboard_tasks.png", "Dashboard – mostrando tareas pendientes"),
    ("04_task_list.png",       "Lista de tareas – vista completa con filtros"),
    ("05_task_detail.png",     "Detalle de tarea – información completa"),
    ("06_focus_page.png",      "Modo Enfoque – selección de tarea o sesión libre"),
    ("07_focus_timer.png",     "Temporizador Pomodoro – en marcha con arco animado"),
    ("08_stats_page.png",      "Estadísticas – KPIs y gráficos calculados desde Realm"),
    ("09_profile_dark.png",    "Perfil – modo oscuro activo con selector de acento"),
    ("10_profile_light.png",   "Perfil – modo claro activo (mismo selector de acento)"),
    ("11_accent_change.png",   "Cambio de color de acento – azul aplicado a toda la app"),
]

for i in range(0, len(screens) - 1, 2):
    f1, c1 = screens[i]
    f2, c2 = screens[i + 1]
    add_two_screenshots(doc, f1, c1, f2, c2)

# última si número impar
if len(screens) % 2 != 0:
    f, c = screens[-1]
    add_screenshot(doc, f, c)

doc.add_page_break()

# ── 6. NAVEGACIÓN ─────────────────────────────────────────
add_heading(doc, "6. Evidencias de navegación")
add_para(doc,
    "La navegación se implementa con GoRouter y StatefulShellRoute. Cada pestaña mantiene su propio "
    "estado de navegación (back-stack independiente). Las rutas anidadas permiten navegar a pantallas "
    "de detalle sin perder el estado de la pestaña padre.",
    size=11)
doc.add_paragraph()

add_para(doc, "Flujo de navegación principal:", bold=True, size=11)
flujos = [
    "Inicio → Dashboard (/dashboard)  [pestaña 1]",
    "Dashboard → Lista de tareas (/tasks)  [pestaña 2]",
    "Lista de tareas → Crear tarea (/tasks/new)  [ruta anidada]",
    "Lista de tareas → Detalle (/tasks/:id)  [ruta anidada]",
    "Detalle → Marcar como completada → volver a lista (pop)",
    "Enfoque (/focus) → Temporizador (/focus/timer?taskId=...)  [ruta anidada]",
    "Perfil (/profile) → Cambio de tema (sin navegación, in-place)",
]
for f in flujos:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f).font.size = Pt(11)

doc.add_page_break()

# ── 7. ACCESO A DATOS (REALM) ─────────────────────────────
add_heading(doc, "7. Evidencias de acceso a datos – Realm")
add_para(doc,
    "La aplicación persiste y recupera todos los datos usando Realm como base de datos local embebida. "
    "No existe ningún dato simulado, hardcodeado ni semilla (seed) de datos.",
    size=11)
doc.add_paragraph()

add_para(doc, "Configuración de Realm (injection_container.dart):", bold=True, size=11)
add_code_block(doc, """\
final config = Configuration.local(
  [TaskObject.schema],
  schemaVersion: 1,
  shouldDeleteIfMigrationNeeded: true,
);
final realm = Realm(config);
sl.registerSingleton<Realm>(realm);
sl.registerSingleton<TaskRealmDataSource>(
    TaskRealmDataSource(sl<Realm>()));""")

doc.add_paragraph()
add_para(doc,
    "El ciclo de vida completo de una tarea en Realm:",
    bold=True, size=11)
ciclo = [
    "Creación: TaskRealmDataSource.createTask() → realm.write(() => realm.add(TaskObject(...)))",
    "Consulta: realm.all<TaskObject>() → mapeado a entidades de dominio",
    "Toggle: realm.write(() { obj.isDone = !obj.isDone; obj.completedAt = ... })",
    "Eliminación: realm.write(() => realm.delete(obj))",
    "Persistencia: Realm escribe en archivo .realm en el directorio de datos de la app",
]
for c in ciclo:
    p = doc.add_paragraph(style="List Number")
    p.add_run(c).font.size = Pt(11)

doc.add_page_break()

# ── 8. CONCLUSIONES ───────────────────────────────────────
add_heading(doc, "8. Conclusiones y trabajo futuro")
add_para(doc,
    "En esta primera entrega se ha implementado una base funcional real y completa de la aplicación KAIROS. "
    "La arquitectura limpia permite añadir nuevas capas (API remota, autenticación, sincronización) "
    "sin modificar la lógica de dominio. Todos los datos se persisten de forma real mediante Realm.",
    size=11)
doc.add_paragraph()
add_para(doc, "Para las siguientes entregas está previsto:", bold=True, size=11)
futuro = [
    "Integración con backend remoto (API REST / Firebase)",
    "Autenticación de usuario (registro e inicio de sesión)",
    "Sincronización offline-first: local primero, sync en background",
    "Notificaciones push para recordatorios de tareas",
    "Integración de IA para optimización automática de agenda",
    "Widget de escritorio / notificación persistente de tarea en curso",
]
for f in futuro:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f).font.size = Pt(11)

doc.add_paragraph()
add_para(doc,
    f"Versión de compilación: KAIROS 2.0.1 · BUILD 2026.04.27 · ©IML",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── GUARDAR ───────────────────────────────────────────────
doc.save(str(OUTPUT))
print(f"Documento generado: {OUTPUT}")
print(f"Pantallas disponibles en {SCREENS_DIR}:")
for f in sorted(SCREENS_DIR.glob("*.png")):
    print(f"  {f.name}  ({f.stat().st_size // 1024} KB)")
