#!/usr/bin/env python3
"""
Generador de DOCX para Entrega 2 - KAIROS
Evidencias de funcionalidades, código y flujo con screenshots
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCREENS_DIR = Path(r"C:\Users\Ismael\Desktop\KAIROS\entregas\entrega2\screens")
OUTPUT = Path(r"C:\Users\Ismael\Desktop\KAIROS\entregas\entrega2\Entrega2_KAIROS.docx")

# ──────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def add_code_block(doc, code, lang="dart"):
    style_name = "Code Block"
    if style_name not in [s.name for s in doc.styles]:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(9)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Cm(0.5)

    for line in code.split("\n"):
        p = doc.add_paragraph(style=style_name)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)

def add_screenshot(doc, filename, caption=""):
    path = SCREENS_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(3.2))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
    else:
        add_para(doc, f"[Captura no disponible: {filename}]", italic=True, color=(180, 0, 0))

def _cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _row_height(row, twips):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(twips))
    trHeight.set(qn('w:type'), 'exact')
    trPr.append(trHeight)

# ──────────────────────────────────────────────────────────
# DOCUMENT
# ──────────────────────────────────────────────────────────

def create_document():
    doc = Document()

    # ─── COVER ───
    cover = doc.add_table(rows=1, cols=1)
    cover.autofit = False
    cover.allow_autofit = False
    cell = cover.rows[0].cells[0]
    _cell_shading(cell, "050505")
    _row_height(cover.rows[0], 9000)

    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("\n\n\nKAIROS 2.0\n")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(240, 230, 215)

    run = p.add_run("Aplicación de Productivity con\nOffline-First & Deep Work\n\n")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(200, 200, 200)

    run = p.add_run("Entrega 2 - Aplicación Funcional\n\n\n")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(180, 180, 180)

    run = p.add_run("Ismael Manzano León\nismaelmanzanoleon@gmail.com\n")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(160, 160, 160)

    run = p.add_run("\nMayo 2026 - v2.0.0")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(140, 140, 140)

    doc.add_page_break()

    # ─── CONTENIDO ───
    add_heading(doc, "KAIROS 2.0", level=1)
    add_para(doc, "Aplicación de Productivity con Offline-First & Deep Work")

    # Identificación
    add_heading(doc, "Identificación del Alumnado", level=2)
    add_para(doc, "Nombre: Ismael Manzano León")
    add_para(doc, "Email: ismaelmanzanoleon@gmail.com")
    add_para(doc, "Fecha: Mayo 2026")
    add_para(doc, "Versión: 2.0.0")
    add_para(doc, "Entrega: 2 (Aplicación Funcional)")

    doc.add_paragraph()  # Espacio

    # Descripción
    add_heading(doc, "Descripción del Proyecto", level=2)
    add_para(doc, "KAIROS es una aplicación Flutter de productivity basada en:")
    add_para(doc, "• Offline-First: Realm como fuente de verdad local, Supabase como remoto", size=10)
    add_para(doc, "• Deep Work: Pomodoro timer con tracking de sesiones dinámicas", size=10)
    add_para(doc, "• Smart Scheduling: Algoritmo heurístico para optimizar orden de tareas", size=10)
    add_para(doc, "• Visual: Dark Glassmorphism theme con warm ivory accents", size=10)

    doc.add_paragraph()

    # Flujo completo
    add_heading(doc, "Flujo Completo de la Aplicación", level=2)
    add_para(doc, "1. Splash → Animación de carga (1.8s, 'realm · syncing local store...')")
    add_para(doc, "2. Onboarding → 3 slides: Offline-First, Smart Scheduling, Deep Work")
    add_para(doc, "3. Login → Email + contraseña con validaciones")
    add_para(doc, "4. Dashboard → Resumen de tareas hoy + energy bar")
    add_para(doc, "5. Task Management → CRUD completo con filtros y swipe actions")
    add_para(doc, "6. Focus Mode → Timer Pomodoro 25min con progreso visual")
    add_para(doc, "7. Estadísticas → KPIs, gráficos 7 días y heatmap 4 semanas")
    add_para(doc, "8. Perfil → Sincronización, temas y configuración")

    doc.add_page_break()

    # SCREENSHOTS - Flujo principal
    add_heading(doc, "Evidencia Visual - Flujo Principal", level=2)

    add_heading(doc, "Splash Screen", level=3)
    add_screenshot(doc, "01_splash.png", "Animación de carga inicial (1.8s)")

    add_heading(doc, "Onboarding", level=3)
    add_screenshot(doc, "02_onboarding_1.png", "Slide 1: Offline-First")
    add_screenshot(doc, "03_onboarding_2.png", "Slide 2: Smart Scheduling")
    add_screenshot(doc, "04_onboarding_3.png", "Slide 3: Deep Work")

    add_heading(doc, "Login", level=3)
    add_screenshot(doc, "05_login.png", "Página de login con validaciones")

    add_heading(doc, "Dashboard", level=3)
    add_screenshot(doc, "06_dashboard.png", "Resumen de tareas y energy bar")

    doc.add_page_break()

    # Task Management
    add_heading(doc, "Gestión de Tareas (CRUD)", level=2)

    add_heading(doc, "Lista de Tareas", level=3)
    add_screenshot(doc, "07_task_list_empty.png", "Lista vacía")
    add_screenshot(doc, "08_task_list.png", "Con tareas")
    add_screenshot(doc, "09_task_filters.png", "Filtros activos")

    add_heading(doc, "Crear Tarea", level=3)
    add_screenshot(doc, "10_create_task.png", "Formulario de creación")

    add_heading(doc, "Detalle de Tarea", level=3)
    add_screenshot(doc, "11_task_detail.png", "Vista de detalle con opción Focus")

    doc.add_page_break()

    # Focus Mode
    add_heading(doc, "Focus Mode - Deep Work", level=2)

    add_heading(doc, "Selector de Tarea", level=3)
    add_screenshot(doc, "12_focus_landing.png", "Seleccionar tarea para enfocar")

    add_heading(doc, "Timer Pomodoro", level=3)
    add_screenshot(doc, "13_focus_timer_start.png", "Timer en inicio (25:00)")
    add_screenshot(doc, "14_focus_timer_running.png", "Timer en ejecución con arc animado")
    add_screenshot(doc, "15_focus_timer_pause.png", "Timer en pausa")
    add_screenshot(doc, "16_focus_complete.png", "Sesión completada")

    doc.add_page_break()

    # Estadísticas
    add_heading(doc, "Estadísticas y Análisis", level=2)
    add_screenshot(doc, "17_stats_overview.png", "KPIs: completadas, tiempo, racha")
    add_screenshot(doc, "18_stats_chart.png", "Gráfico 7 días")
    add_screenshot(doc, "19_stats_heatmap.png", "Heatmap 4 semanas")

    doc.add_page_break()

    # Perfil & Sincronización
    add_heading(doc, "Perfil y Sincronización", level=2)
    add_screenshot(doc, "20_profile.png", "Perfil modo oscuro")
    add_screenshot(doc, "21_profile_light.png", "Perfil modo claro")
    add_screenshot(doc, "22_sync_sheet.png", "Sincronización en progreso")
    add_screenshot(doc, "23_conflict_resolution.png", "Resolución de conflictos")

    doc.add_page_break()

    # Backend & Integración
    add_heading(doc, "Evidencia de Integración Backend", level=2)

    add_heading(doc, "Supabase Setup", level=3)
    add_para(doc, "URL: https://mxhyuzucjygdjmamtcjq.supabase.co", italic=True)
    add_para(doc, "Tabla: tasks (id, title, priority, energy, project, is_completed, is_synced)")
    add_para(doc, "Sincronización: SupabaseSyncService.pushTasks() → upsert en Supabase")
    add_para(doc, "Flujo: Realm local → Detectar cambios → Push a Supabase → Sincronizar")

    add_heading(doc, "Operaciones Reales", level=3)
    add_para(doc, "✅ Crear tarea en Realm (local)")
    add_para(doc, "✅ Modificar tarea local")
    add_para(doc, "✅ Marcar completada (local + sync)")
    add_para(doc, "✅ Perfil: 'Forzar sincronización' con 4 pasos reales")
    add_para(doc, "✅ Datos persistentes en Realm")
    add_para(doc, "✅ Estadísticas calculadas en tiempo real")

    doc.add_page_break()

    # Código Clave
    add_heading(doc, "Fragmentos de Código Clave", level=2)

    add_heading(doc, "FocusBloc con Rondas Dinámicas", level=3)
    add_code_block(doc, """class FocusBloc extends Bloc<FocusEvent, FocusState> {
  int _currentRound = 1;

  void _onTick(FocusTick event, Emitter<FocusState> emit) {
    if (state is FocusRunning) {
      final current = state as FocusRunning;
      if (current.secondsLeft <= 1) {
        _timer?.cancel();
        _currentRound++;
        emit(FocusCompleted(round: current.round));
      } else {
        emit(FocusRunning(
          secondsLeft: current.secondsLeft - 1,
          task: current.task,
          round: current.round
        ));
      }
    }
  }
}""")

    doc.add_paragraph()

    add_heading(doc, "Validaciones - Login", level=3)
    add_code_block(doc, """String? _validateEmail(String? value) {
  if (value == null || value.isEmpty) {
    return 'El correo es obligatorio';
  }
  final emailRegex = RegExp(r'^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$');
  if (!emailRegex.hasMatch(value)) {
    return 'Correo inválido';
  }
  return null;
}

String? _validatePassword(String? value) {
  if (value == null || value.isEmpty) {
    return 'La contraseña es obligatoria';
  }
  if (value.length < 6) {
    return 'Mínimo 6 caracteres';
  }
  return null;
}""")

    doc.add_paragraph()

    add_heading(doc, "SupabaseSyncService", level=3)
    add_code_block(doc, """class SupabaseSyncService {
  Future<int> pushTasks(List<TaskObject> tasks) async {
    int syncedCount = 0;
    for (final task in tasks) {
      if (!task.isSynced) {
        await supabase.from('tasks').upsert({
          'id': task.id.toString(),
          'title': task.title,
          'priority': task.priority,
          'energy': task.energyLevel,
          'project': task.project,
          'is_completed': task.isDone,
          'completed_at': task.completedAt?.toIso8601String(),
          'is_synced': true,
        });
        syncedCount++;
      }
    }
    return syncedCount;
  }
}""")

    doc.add_page_break()

    # Stack Técnico
    add_heading(doc, "Stack Técnico", level=2)

    table = doc.add_table(rows=11, cols=2)
    table.style = 'Light Grid Accent 1'

    headers = table.rows[0].cells
    headers[0].text = "Componente"
    headers[1].text = "Tecnología"

    data = [
        ("Frontend", "Flutter 3.3+, Dart"),
        ("UI Framework", "Material Design 3"),
        ("State Management", "flutter_bloc + Cubit"),
        ("Routing", "go_router 14.0"),
        ("Local Database", "Realm 3.0"),
        ("Remote Backend", "Supabase PostgreSQL"),
        ("HTTP Client", "Dio 5.4"),
        ("DI Container", "get_it 7.7"),
        ("Persistence", "SharedPreferences 2.3"),
        ("UI Toolkit", "Google Fonts, Custom Painters"),
    ]

    for i, (component, tech) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = component
        row_cells[1].text = tech

    doc.add_page_break()

    # Validaciones
    add_heading(doc, "Validaciones Implementadas", level=2)

    add_heading(doc, "Login", level=3)
    add_para(doc, "✅ Email: formato válido (regex validation)")
    add_para(doc, "✅ Contraseña: mínimo 6 caracteres")
    add_para(doc, "✅ Campos obligatorios")

    add_heading(doc, "Create Task", level=3)
    add_para(doc, "✅ Título: obligatorio")
    add_para(doc, "✅ Prioridad: seleccionar de opciones")
    add_para(doc, "✅ Energía: slider 1-5")
    add_para(doc, "✅ Proyecto: opcional")
    add_para(doc, "✅ Fecha: date picker")

    add_heading(doc, "Focus", level=3)
    add_para(doc, "✅ Timer: 25 minutos (Pomodoro)")
    add_para(doc, "✅ Rondas: contador dinámico 1/4 → 2/4 → 3/4 → 4/4")
    add_para(doc, "✅ Pausa/Resume: funcionan correctamente")
    add_para(doc, "✅ Reset: reinicia timer y ronda")

    doc.add_page_break()

    # Glassmorphism
    add_heading(doc, "Mejoras Visuales - Glassmorphism", level=2)

    add_heading(doc, "Paleta de Colores", level=3)
    add_para(doc, "Fondo: #050505 (negro profundo)")
    add_para(doc, "Accent Primario: #F0E6D7 (warm ivory)")
    add_para(doc, "Glass surfaces: rgba(255,255,255,0.03) con blur(16)")
    add_para(doc, "Borders: rgba(255,255,255,0.15)")
    add_para(doc, "Glows: Cool (azul) top-left + Warm (marfil) bottom-right")

    add_heading(doc, "Componentes", level=3)
    add_para(doc, "✅ GlassCard: backdrop-filter blur + inner shadow")
    add_para(doc, "✅ Nav bar: pill flotante con glassmorphism")
    add_para(doc, "✅ FAB: glassmórfico con glow cálido")
    add_para(doc, "✅ Energy bar: gradiente bicolor")
    add_para(doc, "✅ Cards: semi-transparent con borders sutil")
    add_para(doc, "✅ Animaciones suaves: transiciones entre pantallas")

    doc.add_page_break()

    # Estado del Proyecto
    add_heading(doc, "Estado del Proyecto - Entrega 2", level=2)

    add_heading(doc, "Completado ✅", level=3)
    add_para(doc, "✅ Backend Supabase integrado (PostgreSQL remoto)")
    add_para(doc, "✅ Realm offline-first (fuente de verdad local)")
    add_para(doc, "✅ SupabaseSyncService con push() operativo")
    add_para(doc, "✅ isSynced flag para rastrear estado de datos")
    add_para(doc, "✅ Login page con validaciones (email regex, min 6 chars)")
    add_para(doc, "✅ CRUD completo: Crear, Leer, Actualizar, Eliminar tareas")
    add_para(doc, "✅ Filtros: Todas, Pendientes, Completadas, Alta prioridad")
    add_para(doc, "✅ Swipe actions: derecha completa, izquierda elimina")
    add_para(doc, "✅ Focus Mode: Selector de tarea + Timer Pomodoro 25min")
    add_para(doc, "✅ Contador de rondas dinámico (n/4)")
    add_para(doc, "✅ Arc animado mostrando progreso en tiempo real")
    add_para(doc, "✅ Estadísticas: KPIs, gráficos 7 días, heatmap 4 semanas")
    add_para(doc, "✅ Perfil: toggle tema oscuro/claro, selector color acento")
    add_para(doc, "✅ Sincronización: SyncSheet con 4 pasos reales")
    add_para(doc, "✅ Resolución de conflictos versión (ConflictSheet)")
    add_para(doc, "✅ Toggle sync online/offline con animación")
    add_para(doc, "✅ Cerrar sesión funcional")
    add_para(doc, "✅ Dark Glassmorphism con glow ambient")
    add_para(doc, "✅ Todos los 9 requisitos implementados")

    add_heading(doc, "Próximas Entregas", level=3)
    add_para(doc, "[ ] Entrega 3: Autenticación real (Firebase Auth)")
    add_para(doc, "[ ] Algoritmo IA para optimización de tareas")
    add_para(doc, "[ ] Notificaciones push (local + remote)")
    add_para(doc, "[ ] Sistema de puntos/gamificación")

    doc.add_page_break()

    # Conclusión
    add_heading(doc, "Conclusión", level=2)
    add_para(doc, "KAIROS 2.0 es una aplicación de productivity completamente funcional con arquitectura offline-first. Implementa sincronización bidireccional con Supabase, validaciones robustas, UI modern con glassmorphism, y funcionalidad completa de Focus Mode con Pomodoro timer dinámico.")
    add_para(doc, "El proyecto demuestra integración exitosa de:")
    add_para(doc, "• Base de datos local (Realm) y remota (Supabase PostgreSQL)", size=10)
    add_para(doc, "• State management avanzado con BLoC pattern", size=10)
    add_para(doc, "• Custom UI con painters y animaciones", size=10)
    add_para(doc, "• Validaciones y manejo de errores", size=10)
    add_para(doc, "• Routing dinámico con go_router", size=10)
    add_para(doc, "Todas las funcionalidades requeridas para Entrega 2 están implementadas y operativas.", size=10, bold=True)

    return doc

# ──────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generando documento Entrega 2...")
    doc = create_document()

    # Crear directorio si no existe
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    # Guardar
    doc.save(str(OUTPUT))
    print(f"✅ Documento generado: {OUTPUT}")
    print(f"📊 Tamaño: {OUTPUT.stat().st_size / 1024:.1f} KB")
