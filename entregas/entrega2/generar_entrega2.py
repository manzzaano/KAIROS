"""
Genera el DOCX de Entrega 2 - KAIROS App
Requiere: python-docx, Pillow
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCREENS_DIR = Path(r"C:\Users\Ismael\Desktop\KAIROS\entregas\entrega2\screens")
OUTPUT = Path(r"C:\Users\Ismael\Desktop\KAIROS\entregas\entrega2\Entrega2_KAIROS.docx")

# ──────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def add_code_block(doc, code, lang="dart"):
    style_name = "Code Block"
    if style_name not in [s.name for s in doc.styles]:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(9)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Cm(0.5)

    for line in code.split("\n"):
        p = doc.add_paragraph(style=style_name)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_screenshot(doc, filename, caption=""):
    path = SCREENS_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(3.2))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
    else:
        add_para(doc, f"[Captura no disponible: {filename}]", italic=True, color=(180, 0, 0))

def add_toc(doc):
    """Inserta campo TOC de Word con hipervínculos. Requiere Ctrl+A → F9 en Word para actualizar."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_end)


# ── Cover page helpers ────────────────────────────────────

def _cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _row_height(row, twips):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(twips))
    h.set(qn('w:hRule'), 'exact')
    trPr.append(h)

def _no_borders(tbl):
    tbl_elt = tbl._tbl
    tblPr = tbl_elt.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_elt.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def _cell_no_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _cpara(cell, text='', size=11, bold=False, italic=False,
           color=None, align=WD_ALIGN_PARAGRAPH.CENTER,
           sb=0, sa=0, font='Calibri'):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        r.font.name  = font
        r.font.size  = Pt(size)
        r.bold       = bold
        r.italic     = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p

def _insert_page_field(run):
    for tag in ('begin',):
        fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), tag)
        run._r.append(fc)
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    run._r.append(instr)
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run._r.append(fc2)

def build_portada(doc):
    """Portada moderna de 3 filas: fondo oscuro | franja naranja | sección blanca."""
    # Full content width: 21cm - 2*2.5cm = 16cm = 9072 twips
    tbl = doc.add_table(rows=3, cols=1)
    tbl.style = 'Table Grid'
    _no_borders(tbl)

    tbl_elt2 = tbl._tbl
    tblPr = tbl_elt2.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_elt2.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9072')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Heights (A4 content ~24.7cm = 13985 twips)
    _row_height(tbl.rows[0], 6000)   # dark   ~10.6 cm
    _row_height(tbl.rows[1], 220)    # stripe ~0.39 cm
    _row_height(tbl.rows[2], 7765)   # white  ~13.7 cm

    # ── Row 0: dark ──────────────────────────────────────
    dc = tbl.rows[0].cells[0]
    _cell_shading(dc, '0A0A0A')
    _cell_no_borders(dc)
    dc.paragraphs[0].clear()

    _cpara(dc, sb=50, sa=0)          # top spacer
    _cpara(dc, sb=18, sa=0)
    _cpara(dc, sb=18, sa=0)

    # KAIROS title
    p_k = dc.add_paragraph()
    p_k.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_k.paragraph_format.space_before = Pt(6)
    p_k.paragraph_format.space_after  = Pt(0)
    rk = p_k.add_run("KAIROS")
    rk.font.name = 'Calibri'
    rk.font.size = Pt(64)
    rk.bold = True
    rk.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Orange rule below title
    p_rule = dc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule.paragraph_format.space_before = Pt(6)
    p_rule.paragraph_format.space_after  = Pt(10)
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '18')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'FB923C')
    pBdr.append(bot)
    pPr.append(pBdr)

    _cpara(dc, 'Aplicación de gestión de tareas y productividad personal',
           size=12, color=(0xCC, 0xCC, 0xCC), sb=4, sa=6)
    _cpara(dc, 'Una tarea. Un cronómetro. Sin distracciones.',
           size=10, italic=True, color=(0xFB, 0x92, 0x3C), sb=0, sa=0)

    # ── Row 1: orange stripe ──────────────────────────────
    oc = tbl.rows[1].cells[0]
    _cell_shading(oc, 'FB923C')
    _cell_no_borders(oc)
    oc.paragraphs[0].clear()

    # ── Row 2: white info ─────────────────────────────────
    wc = tbl.rows[2].cells[0]
    _cell_shading(wc, 'FFFFFF')
    _cell_no_borders(wc)
    wc.paragraphs[0].clear()

    _cpara(wc, sb=14, sa=0)

    # "ENTREGA 2" badge
    _cpara(wc, 'ENTREGA 2', size=24, bold=True,
           color=(0xFB, 0x92, 0x3C), sb=6, sa=2)
    _cpara(wc, 'Aplicación funcional  ·  40%', size=11,
           color=(0x80, 0x80, 0x80), sb=0, sa=10)

    # Horizontal rule
    p_hr = wc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(0)
    p_hr.paragraph_format.space_after  = Pt(8)
    hrPr = p_hr._p.get_or_add_pPr()
    hrBdr = OxmlElement('w:pBdr')
    hrBot = OxmlElement('w:bottom')
    hrBot.set(qn('w:val'), 'single')
    hrBot.set(qn('w:sz'), '4')
    hrBot.set(qn('w:space'), '1')
    hrBot.set(qn('w:color'), 'E0E0E0')
    hrBdr.append(hrBot)
    hrPr.append(hrBdr)

    # Info table (nested, 2 cols, centered)
    info_data = [
        ("Alumno",               "Ismael Manzano León"),
        ("Ciclo Formativo",      "[Nombre Ciclo Formativo]"),
        ("Centro",               "[Nombre del Centro Educativo]"),
        ("Módulo / Asignatura",  "[Nombre del Módulo]"),
        ("Profesor/a",           "[Nombre del Profesor/a]"),
        ("Fecha de entrega",     "11 de mayo de 2026"),
        ("Versión",              "2.0.0"),
    ]
    itbl = wc.add_table(rows=len(info_data), cols=2)
    itbl.style = 'Table Grid'
    _no_borders(itbl)

    itbl_elt = itbl._tbl
    itblPr = itbl_elt.find(qn('w:tblPr'))
    if itblPr is None:
        itblPr = OxmlElement('w:tblPr')
        itbl_elt.insert(0, itblPr)
    itblW = OxmlElement('w:tblW')
    itblW.set(qn('w:w'), '7200')
    itblW.set(qn('w:type'), 'dxa')
    itblPr.append(itblW)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    itblPr.append(jc)

    for i, (label, value) in enumerate(info_data):
        lc = itbl.rows[i].cells[0]
        vc = itbl.rows[i].cells[1]
        bg = 'F5F5F5' if i % 2 == 0 else 'FAFAFA'
        _cell_shading(lc, bg)
        _cell_shading(vc, bg)
        _cell_no_borders(lc)
        _cell_no_borders(vc)

        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after  = Pt(4)
        lr = lp.add_run(label)
        lr.font.name = 'Calibri'
        lr.font.size = Pt(9)
        lr.bold = True
        lr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vp.paragraph_format.space_before = Pt(4)
        vp.paragraph_format.space_after  = Pt(4)
        vr = vp.add_run(f"  {value}")
        vr.font.name = 'Calibri'
        vr.font.size = Pt(9)
        vr.bold = (i == 0)
        vr.font.color.rgb = RGBColor(0x15, 0x15, 0x15)

    _cpara(wc, sb=10, sa=0)
    _cpara(wc, 'KAIROS 2.0.0  ·  BUILD 2026.05.11  ·  ©IML',
           size=7, color=(0xBB, 0xBB, 0xBB), sb=6, sa=0, font='Courier New')


def apply_headers_footers(doc):
    """Cabecera con nombre del proyecto + alumno. Pie con número de página."""
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # Header
        hdr = section.header
        hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        rl = hp.add_run("KAIROS · Entrega 2 – Aplicación funcional")
        rl.font.size = Pt(9)
        rl.font.name = 'Calibri'
        rl.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        hp.add_run("\t")

        rr = hp.add_run("Ismael Manzano León")
        rr.font.size = Pt(9)
        rr.font.name = 'Calibri'
        rr.bold = True
        rr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        # Right-aligned tab stop at right margin (~9 cm = 5103 twips from left)
        pPr = hp._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9072')
        tabs.append(tab)
        pPr.append(tabs)

        # Thin bottom rule on header
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'DDDDDD')
        pBdr.append(bot)
        pPr.append(pBdr)

        # Footer
        ftr = section.footer
        ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for txt in ('— ', None, ' —'):
            r = fp.add_run(txt or '')
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
            if txt is None:
                _insert_page_field(r)


def add_two_screenshots(doc, file1, cap1, file2, cap2):
    """Intenta poner dos capturas lado a lado en una tabla."""
    p1 = SCREENS_DIR / file1
    p2 = SCREENS_DIR / file2
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    # fila imágenes
    c1 = table.cell(0, 0)
    c2 = table.cell(0, 1)
    for cell, path in ((c1, p1), (c2, p2)):
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            run = para.add_run()
            run.add_picture(str(path), width=Inches(2.8))
        else:
            para.add_run(f"[{path.name} no encontrado]")
    # fila captions
    for cell, cap in ((table.cell(1, 0), cap1), (table.cell(1, 1), cap2)):
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(cap)
        run.italic = True
        run.font.size = Pt(9)
    doc.add_paragraph()


# ──────────────────────────────────────────────────────────
# DOCUMENT
# ──────────────────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── PORTADA ──────────────────────────────────────────────
build_portada(doc)

doc.add_page_break()

# Aplicar headers/footers después de la portada
apply_headers_footers(doc)

# ── ÍNDICE ───────────────────────────────────────────────
add_heading(doc, "Índice")
add_toc(doc)
p_note = doc.add_paragraph()
r_note = p_note.add_run("Nota: abre en Word y pulsa Ctrl+A → F9 para actualizar el índice.")
r_note.italic = True
r_note.font.size = Pt(9)
r_note.font.color.rgb = RGBColor(150, 100, 0)
doc.add_page_break()

# ── 1. DESCRIPCIÓN DEL PROYECTO ──────────────────────────
add_heading(doc, "1. Descripción del proyecto")
add_para(doc,
    "KAIROS es una aplicación móvil de gestión de tareas y productividad personal desarrollada en Flutter. "
    "En esta segunda entrega, la aplicación se ha extendido con sincronización cloud mediante Supabase, "
    "un sistema de onboarding de 3 slides, login con validación de formulario, modo glassmorphism "
    "y una paleta de 17 tokens de color dinámicos. "
    "Implementa Clean Architecture (Domain / Data / Presentation) con BLoC/Cubit como gestor de estado, "
    "Realm como base de datos local offline-first y GoRouter para la navegación.",
    size=11)
doc.add_paragraph()
add_para(doc, "Tecnologías principales:", bold=True, size=11)
tecnologias = [
    "Flutter 3.x (Dart 3)",
    "flutter_bloc / Cubit – gestión de estado reactivo",
    "Realm – base de datos local embebida (offline-first)",
    "go_router – navegación declarativa",
    "get_it – inyección de dependencias",
    "shared_preferences – persistencia de preferencias de usuario",
    "supabase_flutter – sincronización cloud y backend remoto",
    "google_fonts – tipografías personalizadas",
    "glassmorphism – efecto de cristal en widgets UI",
]
for t in tecnologias:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(t).font.size = Pt(11)

doc.add_paragraph()

# ── 2. ARQUITECTURA ──────────────────────────────────────
add_heading(doc, "2. Arquitectura del proyecto")
add_para(doc,
    "El proyecto sigue Clean Architecture dividido en tres capas claramente separadas. "
    "En Entrega 2 se añade una nueva capa de servicios con SupabaseSyncService, "
    "el módulo de onboarding completo y el flujo de autenticación (splash → onboarding → login → dashboard):",
    size=11)
capas = [
    ("Domain layer", "Entidades puras (Task, Priority, EnergyLevel) y contratos de repositorio (ITaskRepository). Sin dependencias de Flutter ni Realm."),
    ("Data layer", "Implementación de repositorios (TaskRepositoryImpl), datasource de Realm (TaskRealmDataSource) y modelos de objeto Realm (TaskObject)."),
    ("Presentation layer", "BLoC (TaskBloc, FocusBloc) y Cubits (ThemeCubit) que emiten estados. Widgets reactivos que escuchan esos estados mediante BlocBuilder."),
    ("Services layer", "SupabaseSyncService para upsert a Supabase cloud. Inicializado en injection_container y registrado con GetIt."),
]
for nombre, desc in capas:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{nombre}: ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

doc.add_paragraph()
add_para(doc, "Estructura de directorios actualizada:", bold=True, size=11)
add_code_block(doc, """\
lib/
├── core/
│   ├── services/supabase_sync_service.dart
│   ├── theme/kairos_colors.dart  (17 tokens)
│   └── router/app_router.dart    (auth flow: splash→onboarding→login)
├── features/
│   ├── onboarding/               (splash, onboarding, login)
│   ├── optimize/                 (OptimizePage)
│   ├── sync/                     (SyncSheet, ConflictSheet)
│   └── ... (igual que E1: tasks, dashboard, focus, stats, profile)
└── shared/widgets/
    ├── glass_card.dart
    ├── kairos_background.dart
    └── offline_banner.dart""", "text")

doc.add_page_break()

# ── 3. FUNCIONALIDADES IMPLEMENTADAS ─────────────────────
add_heading(doc, "3. Funcionalidades implementadas")
add_para(doc,
    "Todas las funcionalidades de Entrega 1 se mantienen. En Entrega 2 se añaden las siguientes:",
    size=11)
doc.add_paragraph()

funcionalidades = [
    ("Splash animado → Onboarding → Login",
     "Flujo de entrada completo: pantalla de splash animada (2s), onboarding con 3 slides "
     "swipeables que presentan la app, y pantalla de login con formulario validado."),
    ("Login con validación de formulario",
     "Validación en tiempo real: email mediante regex (^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$) "
     "y contraseña con mínimo 6 caracteres. GlobalKey<FormState> y .validate(). "
     "Mock auth con delay de 2s → navega a /dashboard."),
    ("Sincronización Supabase (SyncSheet)",
     "Bottom sheet con 4 pasos de sincronización animados. SupabaseSyncService.pushTasks() "
     "hace upsert de tareas no sincronizadas a la tabla tasks de Supabase. "
     "Detección de cambios mediante flag isSynced en cada TaskObject."),
    ("ConflictSheet – resolución de conflictos",
     "Sheet de resolución de conflictos de sincronización. "
     "Permite elegir entre la versión local y la versión remota de una tarea."),
    ("OptimizePage – optimización con IA",
     "Página de optimización de agenda con animaciones de carga que simulan "
     "procesamiento de IA. Reorganiza tareas por prioridad y nivel de energía."),
    ("GlassmorphismUI",
     "GlassCard con efecto glassmorphism, KairosBackground con fondo animado, "
     "OfflineBanner como indicador de modo offline."),
    ("Tema Dark glassmorphism",
     "Modo oscuro con fondo #050505 y warm ivory (#F0E6D7) como accent fijo en dark. "
     "Light mode usa el accent seleccionado por el usuario como color principal, bg #FAFAFA."),
    ("KairosColors – ThemeExtension con 17 tokens",
     "Paleta completa: accent, bg, bg2, bg3, line, line2, text, text2, text3, text4, "
     "glowCool, glowWarm, success, danger, warning, accent2, accentSoft. "
     "Fábricas dark() y light() con tokens completamente diferentes por modo."),
    ("Gestión de tareas (Realm)",
     "CRUD completo de tareas con Realm. Sin datos simulados ni semilla."),
    ("Dashboard, Lista, Detalle, Crear tarea",
     "Pantallas principales de gestión de tareas con navegación completa."),
    ("Modo Enfoque (Pomodoro)",
     "Temporizador de 25 minutos con CustomPainter de arco animado."),
    ("Estadísticas reales",
     "KPIs calculados sobre Realm: racha, gráfico 7 días, heatmap 28 días."),
    ("Tema dinámico (8 acentos + dark/light)",
     "ThemeCubit persiste preferencias en SharedPreferences."),
    ("Navegación real (GoRouter)",
     "StatefulShellRoute con 5 tabs + rutas anidadas + flujo de auth pre-dashboard."),
]

for nombre, desc in funcionalidades:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(f"{nombre}: ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

doc.add_page_break()

# ── 4. FRAGMENTOS DE CÓDIGO RELEVANTES ───────────────────
add_heading(doc, "4. Fragmentos de código relevantes")

# 4.1 SupabaseSyncService
add_heading(doc, "4.1 SupabaseSyncService – sincronización cloud", level=2)
add_para(doc,
    "Servicio de sincronización que hace upsert de tareas locales a Supabase. "
    "Solo sincroniza tareas cuyo flag isSynced sea false.",
    size=11)
add_code_block(doc, """\
// lib/core/services/supabase_sync_service.dart
class SupabaseSyncService {
  final SupabaseClient supabase;

  SupabaseSyncService({required this.supabase});

  Future<int> pushTasks(List<TaskObject> tasks) async {
    int syncedCount = 0;
    for (final task in tasks) {
      if (!task.isSynced) {
        await supabase.from('tasks').upsert({
          'id': task.id.toString(),
          'title': task.title,
          'priority': task.priority,
          'energy': task.energyLevel,
          'project': task.project,
          'is_completed': task.isDone,
          'completed_at': task.completedAt?.toIso8601String(),
          'created_at': task.createdAt?.toIso8601String()
              ?? DateTime.now().toIso8601String(),
          'is_synced': true,
        });
        syncedCount++;
      }
    }
    return syncedCount;
  }
}""")

# 4.2 LoginPage – validación
add_heading(doc, "4.2 LoginPage – validación de formulario", level=2)
add_para(doc,
    "Validación de email con regex y contraseña con longitud mínima. "
    "Mock login con delay de 2 segundos que navega al dashboard.",
    size=11)
add_code_block(doc, """\
// lib/features/onboarding/presentation/pages/login_page.dart

// Validación de email con regex
String? _validateEmail(String? value) {
  if (value == null || value.isEmpty) return 'Introduce tu email';
  final emailRegex = RegExp(r'^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$');
  if (!emailRegex.hasMatch(value)) return 'Email no válido';
  return null;
}

// Validación de contraseña
String? _validatePassword(String? value) {
  if (value == null || value.isEmpty) return 'Introduce tu contraseña';
  if (value.length < 6) return 'Mínimo 6 caracteres';
  return null;
}

// Mock login — 2s delay → context.go('/dashboard')
void _login() {
  if (!_formKey.currentState!.validate()) return;
  setState(() => _loading = true);
  Future.delayed(const Duration(seconds: 2), () {
    if (mounted) {
      setState(() => _loading = false);
      context.go('/dashboard');
    }
  });
}""")

# 4.3 SyncSheet – flujo de sincronización
add_heading(doc, "4.3 SyncSheet – flujo de sincronización con Supabase", level=2)
add_para(doc,
    "Bottom sheet con 4 pasos animados que muestra el progreso de la sincronización. "
    "Llama a SupabaseSyncService para hacer el upsert.",
    size=11)
add_code_block(doc, """\
// lib/features/sync/presentation/widgets/sync_sheet.dart (fragmento)
Future<void> _runSync() async {
  setState(() => _step = 0);
  // Paso 1: conectando con Supabase
  await Future.delayed(const Duration(milliseconds: 600));
  setState(() => _step = 1);
  // Paso 2: leyendo tareas locales
  final tasks = getIt<Realm>().all<TaskObject>().toList();
  await Future.delayed(const Duration(milliseconds: 400));
  setState(() => _step = 2);
  // Paso 3: sincronizando cambios
  final synced = await getIt<SupabaseSyncService>().pushTasks(tasks);
  setState(() => _step = 3);
  // Paso 4: completado
  await Future.delayed(const Duration(milliseconds: 300));
  setState(() {
    _step = 4;
    _syncedCount = synced;
    _done = true;
  });
}""")

# 4.4 KairosColors – ThemeExtension
add_heading(doc, "4.4 KairosColors – ThemeExtension con 17 tokens", level=2)
add_para(doc,
    "Extensión de tema con 17 tokens de color dinámicos. "
    "Dark mode: bg #050505, warm ivory #F0E6D7 como accent fijo. "
    "Light mode: bg #FAFAFA, accent del usuario como color principal.",
    size=11)
add_code_block(doc, """\
// lib/core/theme/kairos_colors.dart
class KairosColors extends ThemeExtension<KairosColors> {
  final Color accent;
  final Color bg, bg2, bg3;
  final Color line, line2;
  final Color text, text2, text3, text4;
  final Color glowCool, glowWarm;
  final Color success, danger, warning;
  final Color accent2, accentSoft;

  const KairosColors({
    required this.accent, required this.bg, required this.bg2,
    required this.bg3, required this.line, required this.line2,
    required this.text, required this.text2, required this.text3,
    required this.text4, required this.glowCool, required this.glowWarm,
    required this.success, required this.danger, required this.warning,
    required this.accent2, required this.accentSoft,
  });

  factory KairosColors.dark({Color? accent}) => KairosColors(
    accent: const Color(0xFFF0E6D7),   // warm ivory fijo en dark
    bg: const Color(0xFF050505),
    bg2: const Color(0xFF101010),
    bg3: const Color(0xFF1A1A1A),
    line: const Color(0xFF2A2A2A),
    line2: const Color(0xFF383838),
    text: const Color(0xFFF0E6D7),
    text2: const Color(0xFFB8A89A),
    text3: const Color(0xFF7A6A60),
    text4: const Color(0xFF4A3A30),
    glowCool: const Color(0x2264B5FF),
    glowWarm: const Color(0x22F0E6D7),
    success: const Color(0xFF4ADE80),
    danger: const Color(0xFFF87171),
    warning: const Color(0xFFFBBF24),
    accent2: const Color(0xFFD4C4B0),
    accentSoft: const Color(0x33F0E6D7),
  );

  factory KairosColors.light({required Color accent}) => KairosColors(
    accent: accent,
    bg: const Color(0xFFFAFAFA),
    bg2: const Color(0xFFF0F0F0),
    bg3: const Color(0xFFE8E8E8),
    line: const Color(0xFFDDDDDD),
    line2: const Color(0xFFCCCCCC),
    text: const Color(0xFF111111),
    text2: const Color(0xFF444444),
    text3: const Color(0xFF777777),
    text4: const Color(0xFFAAAAAA),
    glowCool: const Color(0x2264B5FF),
    glowWarm: const Color(0x22FB923C),
    success: const Color(0xFF16A34A),
    danger: const Color(0xFFDC2626),
    warning: const Color(0xFFD97706),
    accent2: accent.withOpacity(0.7),
    accentSoft: accent.withOpacity(0.15),
  );
}""")

# 4.5 AppRouter – flujo de auth
add_heading(doc, "4.5 AppRouter – flujo de autenticación", level=2)
add_para(doc,
    "Router actualizado con rutas de onboarding/auth antes del dashboard. "
    "StatefulShellRoute para 5 tabs y rutas adicionales para optimize y sync.",
    size=11)
add_code_block(doc, """\
// lib/core/router/app_router.dart (fragmento)
final router = GoRouter(
  initialLocation: '/splash',
  routes: [
    GoRoute(path: '/splash',     builder: (_, __) => const SplashPage()),
    GoRoute(path: '/onboarding', builder: (_, __) => const OnboardingPage()),
    GoRoute(path: '/login',      builder: (_, __) => const LoginPage()),
    GoRoute(path: '/optimize',   builder: (_, __) => const OptimizePage()),
    GoRoute(path: '/create-task',builder: (_, __) => const CreateTaskPage()),
    GoRoute(
      path: '/task/:id',
      builder: (_, state) => TaskDetailPage(
        taskId: state.pathParameters['id']!),
    ),
    GoRoute(
      path: '/focus/timer',
      builder: (_, state) => FocusTimerPage(
        taskId: state.uri.queryParameters['taskId']),
    ),
    StatefulShellRoute.indexedStack(
      builder: (_, __, shell) => AppShell(shell: shell),
      branches: [
        StatefulShellBranch(routes: [GoRoute(path: '/dashboard', ...)]),
        StatefulShellBranch(routes: [GoRoute(path: '/tasks', ...)]),
        StatefulShellBranch(routes: [GoRoute(path: '/focus', ...)]),
        StatefulShellBranch(routes: [GoRoute(path: '/stats', ...)]),
        StatefulShellBranch(routes: [GoRoute(path: '/profile', ...)]),
      ],
    ),
  ],
);""")

# 4.6 Inyección de dependencias
add_heading(doc, "4.6 Inyección de dependencias – setupServiceLocator()", level=2)
add_para(doc,
    "Registro completo de dependencias con GetIt. "
    "Supabase se inicializa primero, luego Realm y el resto de servicios.",
    size=11)
add_code_block(doc, """\
// lib/core/di/injection_container.dart
void setupServiceLocator() {
  final supabaseClient = Supabase.instance.client;
  getIt.registerSingleton<SupabaseClient>(supabaseClient);
  getIt.registerSingleton<SupabaseSyncService>(
      SupabaseSyncService(supabase: supabaseClient));

  final config = Configuration.local(
    [TaskObject.schema],
    schemaVersion: 1,
    shouldDeleteIfMigrationNeeded: true,
  );
  final realm = Realm(config);
  getIt.registerSingleton<Realm>(realm);

  getIt.registerSingleton<TaskRealmDataSource>(
      TaskRealmDataSource(getIt<Realm>()));
  getIt.registerSingleton<ITaskRepository>(
      TaskRepositoryImpl(getIt<TaskRealmDataSource>()));

  getIt.registerFactory<GetTasksUseCase>(
      () => GetTasksUseCase(getIt<ITaskRepository>()));
  getIt.registerFactory<CreateTaskUseCase>(
      () => CreateTaskUseCase(getIt<ITaskRepository>()));
  getIt.registerFactory<ToggleTaskUseCase>(
      () => ToggleTaskUseCase(getIt<ITaskRepository>()));
  getIt.registerFactory<DeleteTaskUseCase>(
      () => DeleteTaskUseCase(getIt<ITaskRepository>()));

  getIt.registerFactory<TaskBloc>(() => TaskBloc(
    getTasks:    getIt<GetTasksUseCase>(),
    createTask:  getIt<CreateTaskUseCase>(),
    toggleTask:  getIt<ToggleTaskUseCase>(),
    deleteTask:  getIt<DeleteTaskUseCase>(),
  ));

  getIt.registerSingleton<ThemeCubit>(ThemeCubit());
}""")

doc.add_page_break()

# ── 5. CAPTURAS DE PANTALLA ───────────────────────────────
add_heading(doc, "5. Capturas de pantalla – Aplicación en ejecución real")
add_para(doc,
    "Las siguientes capturas muestran la aplicación KAIROS 2.0 ejecutándose en un emulador Android "
    "(Medium Phone API 35 – x86_64) con el SDK de Flutter compilado en modo debug. "
    "Se incluye el nuevo flujo de onboarding, login, sincronización con Supabase y la UI glassmorphism.",
    size=11)
doc.add_paragraph()

screens = [
    ("01_splash.png",           "Splash – pantalla de carga animada"),
    ("02_onboarding_1.png",     "Onboarding – slide 1: bienvenida"),
    ("03_onboarding_2.png",     "Onboarding – slide 2: funcionalidades"),
    ("04_onboarding_3.png",     "Onboarding – slide 3: sincronización"),
    ("05_login.png",            "Login – formulario con validación"),
    ("06_login_validation.png", "Login – errores de validación activos"),
    ("07_dashboard.png",        "Dashboard – pantalla principal con tareas"),
    ("08_optimize.png",         "Optimize – optimización con IA en curso"),
    ("09_dashboard_tasks.png",  "Dashboard – tareas tras optimize"),
    ("10_create_task.png",      "Crear tarea – formulario completo"),
    ("11_task_list.png",        "Lista de tareas – con filtros"),
    ("12_task_detail.png",      "Detalle de tarea – información completa"),
    ("13_focus_page.png",       "Modo Enfoque – selección de sesión"),
    ("14_focus_timer.png",      "Temporizador Pomodoro – en marcha"),
    ("15_stats_page.png",       "Estadísticas – KPIs y gráficos"),
    ("16_profile_dark.png",     "Perfil – modo oscuro + selector de acento"),
    ("17_sync_sheet.png",       "Sync Sheet – sincronización con Supabase"),
    ("18_accent_change.png",    "Perfil – cambio de color de acento"),
]

for i in range(0, len(screens) - 1, 2):
    f1, c1 = screens[i]
    f2, c2 = screens[i + 1]
    add_two_screenshots(doc, f1, c1, f2, c2)

# última si número impar
if len(screens) % 2 != 0:
    f, c = screens[-1]
    add_screenshot(doc, f, c)

doc.add_page_break()

# ── 6. EVIDENCIAS DE NAVEGACIÓN ───────────────────────────
add_heading(doc, "6. Evidencias de navegación")
add_para(doc,
    "La navegación se implementa con GoRouter. El flujo de autenticación precede al dashboard: "
    "la app arranca en /splash, pasa por /onboarding y /login antes de dar acceso a las 5 tabs. "
    "Cada pestaña mantiene su propio estado de navegación (back-stack independiente).",
    size=11)
doc.add_paragraph()

add_para(doc, "Flujo de navegación completo:", bold=True, size=11)
flujos = [
    "Arranque → Splash (/splash) con animación de 2 segundos",
    "Splash → Onboarding (/onboarding) con 3 slides swipeables",
    "Slide 1: bienvenida a KAIROS · Slide 2: funcionalidades · Slide 3: sync cloud",
    "Onboarding → Login (/login) con formulario de validación",
    "Login (mock, 2s delay) → Dashboard (/dashboard)  [pestaña 1]",
    "Dashboard → Lista de tareas (/tasks)  [pestaña 2]",
    "Lista de tareas → Crear tarea (/create-task)  [ruta separada]",
    "Lista de tareas → Detalle (/task/:id)  [ruta separada]",
    "Dashboard → Enfoque (/focus) → Temporizador (/focus/timer)  [pestaña 3]",
    "Dashboard → Estadísticas (/stats)  [pestaña 4]",
    "Dashboard → Perfil (/profile) → SyncSheet (bottom sheet)  [pestaña 5]",
    "Perfil → OptimizePage (/optimize)  [ruta separada]",
]
for f in flujos:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f).font.size = Pt(11)

doc.add_page_break()

# ── 7. EVIDENCIAS DE BACKEND (SUPABASE) ──────────────────
add_heading(doc, "7. Evidencias de backend – Supabase")
add_para(doc,
    "La aplicación utiliza Supabase como backend cloud. Las tareas se sincronizan mediante upsert "
    "a la tabla tasks cuando el usuario activa la sincronización desde el SyncSheet. "
    "La detección de cambios se realiza mediante el flag isSynced de cada TaskObject en Realm.",
    size=11)
doc.add_paragraph()

add_para(doc, "Schema SQL de la tabla tasks en Supabase:", bold=True, size=11)
add_code_block(doc, """\
-- SQL ejecutado en Supabase Dashboard → SQL Editor
create table tasks (
  id uuid primary key,
  title text not null,
  priority text,
  energy int,
  project text,
  is_completed boolean default false,
  completed_at timestamptz,
  created_at timestamptz default now(),
  is_synced boolean default true
);""", "sql")

doc.add_paragraph()
add_para(doc, "Flujo de sincronización:", bold=True, size=11)
ciclo_sync = [
    "Usuario crea/modifica tarea → Realm escribe localmente → isSynced = false",
    "Usuario abre SyncSheet → _runSync() se ejecuta en 4 pasos animados",
    "SupabaseSyncService.pushTasks() itera tareas con isSynced = false",
    "Para cada tarea no sincronizada: supabase.from('tasks').upsert({...})",
    "Tras upsert exitoso: is_synced = true en Supabase y en Realm",
    "SyncSheet muestra el contador de tareas sincronizadas y estado 'Completado'",
]
for c in ciclo_sync:
    p = doc.add_paragraph(style="List Number")
    p.add_run(c).font.size = Pt(11)

doc.add_paragraph()
add_para(doc, "Inyección de SupabaseClient:", bold=True, size=11)
add_para(doc,
    "Supabase.instance.client se registra en GetIt como singleton. "
    "SupabaseSyncService recibe el cliente por inyección de dependencias y "
    "se registra también como singleton para ser accesible desde cualquier widget.",
    size=11)

doc.add_page_break()

# ── 8. VALIDACIONES Y FLUJO DE USO ───────────────────────
add_heading(doc, "8. Validaciones y flujo de uso")
add_para(doc,
    "La aplicación implementa validación de formularios en la pantalla de login "
    "usando el sistema nativo de Flutter Form/FormField.",
    size=11)
doc.add_paragraph()

add_para(doc, "Reglas de validación del formulario de login:", bold=True, size=11)
validaciones = [
    "Email: regex ^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$ — detecta formato básico nombre@dominio.ext",
    "Contraseña: mínimo 6 caracteres — muestra error si longitud < 6",
    "Validación disparada al pulsar 'Iniciar sesión' mediante _formKey.currentState!.validate()",
    "Errores mostrados inline bajo cada campo (comportamiento nativo de TextFormField)",
]
for v in validaciones:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(v).font.size = Pt(11)

doc.add_paragraph()
add_para(doc, "Flujo de autenticación (mock):", bold=True, size=11)
flujo_auth = [
    "Usuario introduce email y contraseña válidos",
    "Pulsa 'Iniciar sesión' → _login() valida el formulario",
    "Si válido: setState(_loading = true) → muestra CircularProgressIndicator",
    "Future.delayed(2s) → setState(_loading = false) → context.go('/dashboard')",
    "El dashboard carga con todas las tareas de Realm disponibles",
]
for f in flujo_auth:
    p = doc.add_paragraph(style="List Number")
    p.add_run(f).font.size = Pt(11)

doc.add_paragraph()
add_para(doc, "Modo offline:", bold=True, size=11)
add_para(doc,
    "La pantalla de login incluye la opción 'Continuar sin sincronizar' que permite "
    "acceder al dashboard directamente sin pasar por autenticación. "
    "En este modo el OfflineBanner se muestra en el dashboard indicando "
    "que los datos no se sincronizarán con Supabase.",
    size=11)

doc.add_page_break()

# ── 9. CONCLUSIONES ───────────────────────────────────────
add_heading(doc, "9. Conclusiones y trabajo futuro")
add_para(doc,
    "En esta segunda entrega se ha completado la aplicación funcional de KAIROS con "
    "sincronización cloud mediante Supabase, un sistema de onboarding completo, "
    "login con validación real y una interfaz glassmorphism con 17 tokens de color dinámicos. "
    "La arquitectura limpia facilita la extensión hacia las funcionalidades previstas para E3 "
    "sin modificar la lógica de dominio existente.",
    size=11)
doc.add_paragraph()
add_para(doc, "Para la Entrega 3 está previsto:", bold=True, size=11)
futuro = [
    "Autenticación real con Supabase Auth (registro, login, recuperación de contraseña)",
    "IA real para optimización de agenda (integración con API de OpenAI o similar)",
    "Notificaciones push para recordatorios de tareas (flutter_local_notifications)",
    "Sincronización bidireccional: pull de cambios remotos + merge con datos locales",
    "Resolución automática de conflictos (last-write-wins o estrategia configurable)",
    "Widget de escritorio / notificación persistente de tarea en curso",
    "Tests unitarios y de integración con coverage mínimo del 70%",
]
for f in futuro:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f).font.size = Pt(11)

doc.add_paragraph()
add_para(doc,
    "KAIROS 2.0.0  ·  BUILD 2026.05.11  ·  ©IML",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── GUARDAR ───────────────────────────────────────────────
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f"Documento generado: {OUTPUT}")
print(f"Pantallas disponibles en {SCREENS_DIR}:")
if SCREENS_DIR.exists():
    for f in sorted(SCREENS_DIR.glob("*.png")):
        print(f"  {f.name}  ({f.stat().st_size // 1024} KB)")
else:
    print("  (directorio de pantallas no encontrado)")
